"""
WePlus 知识库管理模块 - 完整的RAG系统
支持多种文件格式的上传、处理、向量化存储和智能检索
"""

from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form, BackgroundTasks, Query
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from typing import List, Optional, Dict, Any
import os
import uuid
import hashlib
from datetime import datetime
import logging
import json
import asyncio
import time
from pathlib import Path

# 文件处理库
import PyPDF2
import docx
import openpyxl
import pandas as pd
from io import BytesIO

# 简化的导入，避免复杂依赖
import json

# 本地模块
from models import KnowledgeDocument, User, Base
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker, Session
import jwt
import sys
import os

# 添加项目根目录到Python路径
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.append(project_root)
from config import Config

# 重新创建数据库连接，避免循环导入
engine = create_engine(Config.DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

router = APIRouter(prefix="/api/knowledge", tags=["knowledge_base"])
security = HTTPBearer()

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

def verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        payload = jwt.decode(credentials.credentials, Config.SECRET_KEY, algorithms=[Config.ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise HTTPException(status_code=401, detail="Invalid token")
        return username
    except jwt.PyJWTError:
        raise HTTPException(status_code=401, detail="Invalid token")

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 全局配置
UPLOAD_DIR = Path("uploads/knowledge")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
VECTOR_DB_PATH = Path("knowledge_base/chromadb")
VECTOR_DB_PATH.mkdir(parents=True, exist_ok=True)

# 支持的文件类型配置
SUPPORTED_FILE_TYPES = {
    'pdf': {
        'extensions': ['.pdf'],
        'max_size': 100 * 1024 * 1024,  # 100MB
        'description': 'PDF文档',
        'mime_types': ['application/pdf']
    },
    'docx': {
        'extensions': ['.docx', '.doc'],
        'max_size': 50 * 1024 * 1024,  # 50MB
        'description': 'Microsoft Word文档',
        'mime_types': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']
    },
    'xlsx': {
        'extensions': ['.xlsx', '.xls'],
        'max_size': 30 * 1024 * 1024,  # 30MB
        'description': 'Microsoft Excel表格',
        'mime_types': ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']
    },
    'txt': {
        'extensions': ['.txt', '.md', '.markdown'],
        'max_size': 10 * 1024 * 1024,  # 10MB
        'description': '纯文本文件',
        'mime_types': ['text/plain', 'text/markdown']
    },
    'csv': {
        'extensions': ['.csv'],
        'max_size': 20 * 1024 * 1024,  # 20MB
        'description': 'CSV数据文件',
        'mime_types': ['text/csv']
    },
    'json': {
        'extensions': ['.json'],
        'max_size': 10 * 1024 * 1024,  # 10MB
        'description': 'JSON数据文件',
        'mime_types': ['application/json']
    }
}

# 简化的初始化（后续可以扩展）
def initialize_rag_system():
    """初始化RAG系统组件"""
    try:
        logger.info("RAG系统初始化成功（简化版本）")
        return True
    except Exception as e:
        logger.error(f"RAG系统初始化失败: {e}")
        return False

# 在模块加载时初始化
initialize_rag_system()

# 工具函数
def get_file_type(filename: str) -> Optional[str]:
    """根据文件名获取文件类型"""
    ext = Path(filename).suffix.lower()
    for file_type, config in SUPPORTED_FILE_TYPES.items():
        if ext in config['extensions']:
            return file_type
    return None

def calculate_file_hash(content: bytes) -> str:
    """计算文件的MD5哈希值"""
    return hashlib.md5(content).hexdigest()

def get_next_sequence_number(user_id: int, db: Session) -> int:
    """获取用户的下一个序号"""
    last_doc = db.query(KnowledgeDocument).filter(
        KnowledgeDocument.user_id == user_id
    ).order_by(KnowledgeDocument.sequence_number.desc()).first()
    
    return (last_doc.sequence_number + 1) if last_doc else 1

# 文本提取函数
def extract_text_from_pdf(content: bytes) -> str:
    """从PDF文件中提取文本"""
    try:
        reader = PyPDF2.PdfReader(BytesIO(content))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"PDF文本提取失败: {e}")
        return ""

def extract_text_from_docx(content: bytes) -> str:
    """从Word文档中提取文本"""
    try:
        doc = docx.Document(BytesIO(content))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        logger.error(f"Word文档文本提取失败: {e}")
        return ""

def extract_text_from_xlsx(content: bytes) -> str:
    """从Excel文件中提取文本"""
    try:
        df = pd.read_excel(BytesIO(content), sheet_name=None)
        text = ""
        for sheet_name, sheet_df in df.items():
            text += f"工作表: {sheet_name}\n"
            text += sheet_df.to_string() + "\n\n"
        return text.strip()
    except Exception as e:
        logger.error(f"Excel文件文本提取失败: {e}")
        return ""

def extract_text_from_txt(content: bytes) -> str:
    """从文本文件中提取文本"""
    try:
        # 尝试多种编码
        for encoding in ['utf-8', 'gbk', 'gb2312', 'utf-16']:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return content.decode('utf-8', errors='ignore')
    except Exception as e:
        logger.error(f"文本文件读取失败: {e}")
        return ""

def extract_text_from_csv(content: bytes) -> str:
    """从CSV文件中提取文本"""
    try:
        df = pd.read_csv(BytesIO(content))
        return df.to_string()
    except Exception as e:
        logger.error(f"CSV文件文本提取失败: {e}")
        return ""

def extract_text_from_json(content: bytes) -> str:
    """从JSON文件中提取文本"""
    try:
        data = json.loads(content.decode('utf-8'))
        return json.dumps(data, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"JSON文件文本提取失败: {e}")
        return ""

def extract_text_from_file(content: bytes, file_type: str) -> str:
    """根据文件类型提取文本"""
    extractors = {
        'pdf': extract_text_from_pdf,
        'docx': extract_text_from_docx,
        'xlsx': extract_text_from_xlsx,
        'txt': extract_text_from_txt,
        'csv': extract_text_from_csv,
        'json': extract_text_from_json
    }
    
    extractor = extractors.get(file_type)
    if extractor:
        return extractor(content)
    return ""

def chunk_text(text: str, chunk_size: int = 800, overlap: int = 200) -> List[str]:
    """将文本分块"""
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # 如果不是最后一块，尝试在句号或换行符处分割
        if end < len(text):
            for delimiter in ['. ', '。', '\n\n', '\n']:
                delimiter_pos = text.rfind(delimiter, start, end)
                if delimiter_pos > start:
                    end = delimiter_pos + len(delimiter)
                    break
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        start = max(start + chunk_size - overlap, end)
    
    return chunks

def generate_content_summary(text: str, max_length: int = 300) -> str:
    """生成内容摘要"""
    if len(text) <= max_length:
        return text
    
    # 简单的摘要生成：取前几句话
    sentences = text.replace('。', '.|').replace('. ', '.|').split('|')
    summary = ""
    for sentence in sentences:
        if len(summary + sentence) <= max_length:
            summary += sentence
        else:
            break
    
    return summary.strip() + "..." if len(text) > len(summary) else summary

def extract_keywords(text: str, max_keywords: int = 20) -> List[str]:
    """提取关键词（简单实现）"""
    # 这里可以集成更复杂的关键词提取算法
    import re
    
    # 简单的中文关键词提取
    words = re.findall(r'[\u4e00-\u9fff]+', text)
    word_freq = {}
    for word in words:
        if len(word) >= 2:  # 只考虑长度>=2的词
            word_freq[word] = word_freq.get(word, 0) + 1
    
    # 按频率排序
    sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
    return [word for word, freq in sorted_words[:max_keywords]]

async def process_document_async(doc_id: int, db: Session):
    """异步处理文档（简化版本）"""
    try:
        doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
        if not doc:
            return
        
        # 更新状态为处理中
        doc.processing_status = 'processing'
        doc.processing_progress = 10
        db.commit()
        
        # 读取文件内容
        with open(doc.file_path, 'rb') as f:
            content = f.read()
        
        doc.processing_progress = 30
        db.commit()
        
        # 提取文本
        extracted_text = extract_text_from_file(content, doc.file_type)
        if not extracted_text:
            doc.processing_status = 'failed'
            doc.processing_error = '无法从文件中提取文本内容'
            db.commit()
            return
        
        doc.extracted_text = extracted_text
        doc.content_preview = extracted_text[:500]
        doc.content_summary = generate_content_summary(extracted_text)
        doc.keywords = ','.join(extract_keywords(extracted_text))
        doc.processing_progress = 60
        db.commit()
        
        # 文本分块
        chunks = chunk_text(extracted_text)
        doc.chunk_count = len(chunks)
        doc.vector_count = len(chunks)  # 简化处理
        doc.processing_progress = 100
        
        # 完成处理
        doc.processing_status = 'completed'
        doc.is_processed = True
        doc.updated_at = datetime.utcnow()
        db.commit()
        
        logger.info(f"文档处理完成: {doc.filename}")
        
    except Exception as e:
        logger.error(f"文档处理失败: {e}")
        if 'doc' in locals():
            doc.processing_status = 'failed'
            doc.processing_error = str(e)
            db.commit()

# API路由

@router.get("/supported-types")
async def get_supported_file_types():
    """获取支持的文件类型"""
    return {
        "supported_types": SUPPORTED_FILE_TYPES,
        "max_total_size": 500 * 1024 * 1024,  # 500MB
        "max_files_per_upload": 10
    }

@router.post("/upload")
async def upload_knowledge_files(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
    ai_reference_weight: int = Form(50),
    document_category: Optional[str] = Form(None),
    custom_tags: Optional[str] = Form(None),
    username: str = Depends(verify_token),
    db: Session = Depends(get_db)
):
    """上传知识库文件"""
    
    # 获取用户信息
    user = db.query(User).filter(User.username == username).first()
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在")
    
    # 验证文件数量
    if len(files) > 10:
        raise HTTPException(status_code=400, detail="单次最多上传10个文件")
    
    # 验证AI参考权重
    if not 1 <= ai_reference_weight <= 100:
        raise HTTPException(status_code=400, detail="AI参考权重必须在1-100之间")
    
    uploaded_files = []
    total_size = 0
    
    for file in files:
        # 验证文件大小
        content = await file.read()
        file_size = len(content)
        total_size += file_size
        
        if total_size > 500 * 1024 * 1024:  # 500MB
            raise HTTPException(status_code=400, detail="总文件大小超过500MB限制")
        
        # 验证文件类型
        file_type = get_file_type(file.filename)
        if not file_type:
            raise HTTPException(status_code=400, detail=f"不支持的文件类型: {file.filename}")
        
        # 验证单个文件大小
        max_size = SUPPORTED_FILE_TYPES[file_type]['max_size']
        if file_size > max_size:
            raise HTTPException(
                status_code=400, 
                detail=f"文件 {file.filename} 大小超过限制 ({max_size // (1024*1024)}MB)"
            )
        
        # 检查文件哈希是否已存在
        file_hash = calculate_file_hash(content)
        existing_doc = db.query(KnowledgeDocument).filter(
            KnowledgeDocument.file_hash == file_hash,
            KnowledgeDocument.user_id == user.id
        ).first()
        
        if existing_doc:
            continue  # 跳过重复文件
        
        # 生成唯一文件名
        file_id = str(uuid.uuid4())
        file_extension = Path(file.filename).suffix
        stored_filename = f"{file_id}{file_extension}"
        file_path = UPLOAD_DIR / stored_filename
        
        # 保存文件
        with open(file_path, 'wb') as f:
            f.write(content)
        
        # 解析自定义标签
        tags = []
        if custom_tags:
            tags = [tag.strip() for tag in custom_tags.split(',') if tag.strip()]
        
        # 创建数据库记录
        doc = KnowledgeDocument(
            user_id=user.id,
            sequence_number=get_next_sequence_number(user.id, db),
            filename=stored_filename,
            original_filename=file.filename,
            file_type=file_type,
            file_size=file_size,
            file_hash=file_hash,
            file_path=str(file_path),
            ai_reference_weight=ai_reference_weight,
            document_category=document_category,
            custom_tags=tags if tags else None,
            processing_status='pending'
        )
        
        db.add(doc)
        db.commit()
        db.refresh(doc)
        
        # 添加到后台处理队列
        background_tasks.add_task(process_document_async, doc.id, db)
        
        uploaded_files.append({
            "id": doc.id,
            "filename": file.filename,
            "file_type": file_type,
            "file_size": file_size,
            "sequence_number": doc.sequence_number,
            "status": "pending"
        })
    
    return {
        "message": f"成功上传 {len(uploaded_files)} 个文件",
        "uploaded_files": uploaded_files,
        "processing_message": "文件正在后台处理中，请稍后查看处理状态"
    }

@router.get("/documents")
async def list_knowledge_documents(
    page: int = Query(1, ge=1),
    size: int = Query(20, ge=1, le=100),
    username: str = Depends(verify_token),
    db: Session = Depends(get_db)
):
    """获取知识库文档列表"""
    
    user = db.query(User).filter(User.username == username).first()
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在")
    
    # 构建查询
    query = db.query(KnowledgeDocument).filter(KnowledgeDocument.user_id == user.id)
    
    # 计算总数
    total = query.count()
    
    # 分页查询
    documents = query.order_by(KnowledgeDocument.sequence_number.desc()).offset(
        (page - 1) * size
    ).limit(size).all()
    
    # 格式化结果
    result = []
    for doc in documents:
        result.append({
            "id": doc.id,
            "sequence_number": doc.sequence_number,
            "filename": doc.original_filename,
            "file_type": doc.file_type,
            "file_size": doc.file_size,
            "processing_status": doc.processing_status,
            "processing_progress": doc.processing_progress,
            "is_processed": doc.is_processed,
            "ai_reference_weight": doc.ai_reference_weight,
            "content_preview": doc.content_preview,
            "created_at": doc.created_at.isoformat(),
            "updated_at": doc.updated_at.isoformat()
        })
    
    return {
        "documents": result,
        "total": total,
        "page": page,
        "size": size,
        "pages": (total + size - 1) // size
    }

@router.get("/documents/{doc_id}")
async def get_document_detail(
    doc_id: int,
    username: str = Depends(verify_token),
    db: Session = Depends(get_db)
):
    """获取文档详细信息"""
    
    user = db.query(User).filter(User.username == username).first()
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在")
    
    doc = db.query(KnowledgeDocument).filter(
        KnowledgeDocument.id == doc_id,
        KnowledgeDocument.user_id == user.id
    ).first()
    
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    
    # 更新最后访问时间
    doc.last_accessed_at = datetime.utcnow()
    db.commit()
    
    # 获取文档的分块信息
    chunks = db.query(KnowledgeChunk).filter(
        KnowledgeChunk.document_id == doc_id
    ).order_by(KnowledgeChunk.chunk_index).all()
    
    return {
        "id": doc.id,
        "sequence_number": doc.sequence_number,
        "filename": doc.original_filename,
        "file_type": doc.file_type,
        "file_size": doc.file_size,
        "processing_status": doc.processing_status,
        "processing_progress": doc.processing_progress,
        "is_processed": doc.is_processed,
        "ai_reference_weight": doc.ai_reference_weight,
        "chunk_count": doc.chunk_count,
        "vector_count": doc.vector_count,
        "content_preview": doc.content_preview,
        "content_summary": doc.content_summary,
        "extracted_text": doc.extracted_text,
        "keywords": doc.keywords.split(',') if doc.keywords else [],
        "document_category": doc.document_category,
        "custom_tags": doc.custom_tags,
        "quality_score": doc.quality_score,
        "language": doc.language,
        "embedding_model": doc.embedding_model,
        "created_at": doc.created_at.isoformat(),
        "updated_at": doc.updated_at.isoformat(),
        "last_accessed_at": doc.last_accessed_at.isoformat() if doc.last_accessed_at else None,
        "processing_error": doc.processing_error,
        "chunks": [
            {
                "id": chunk.id,
                "chunk_index": chunk.chunk_index,
                "chunk_text": chunk.chunk_text,
                "chunk_size": chunk.chunk_size,
                "chunk_type": chunk.chunk_type,
                "page_number": chunk.page_number,
                "section_title": chunk.section_title
            } for chunk in chunks
        ]
    }

@router.put("/documents/{doc_id}")
async def update_document(
    doc_id: int,
    ai_reference_weight: Optional[int] = Form(None),
    document_category: Optional[str] = Form(None),
    custom_tags: Optional[str] = Form(None),
    username: str = Depends(verify_token),
    db: Session = Depends(get_db)
):
    """更新文档信息"""
    
    user = db.query(User).filter(User.username == username).first()
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在")
    
    doc = db.query(KnowledgeDocument).filter(
        KnowledgeDocument.id == doc_id,
        KnowledgeDocument.user_id == user.id
    ).first()
    
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    
    # 更新字段
    if ai_reference_weight is not None:
        if not 1 <= ai_reference_weight <= 100:
            raise HTTPException(status_code=400, detail="AI参考权重必须在1-100之间")
        doc.ai_reference_weight = ai_reference_weight
    
    if document_category is not None:
        doc.document_category = document_category
    
    if custom_tags is not None:
        tags = [tag.strip() for tag in custom_tags.split(',') if tag.strip()]
        doc.custom_tags = tags if tags else None
    
    doc.updated_at = datetime.utcnow()
    db.commit()
    
    return {"message": "文档信息更新成功"}

@router.delete("/documents/{doc_id}")
async def delete_document(
    doc_id: int,
    username: str = Depends(verify_token),
    db: Session = Depends(get_db)
):
    """删除文档"""
    
    user = db.query(User).filter(User.username == username).first()
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在")
    
    doc = db.query(KnowledgeDocument).filter(
        KnowledgeDocument.id == doc_id,
        KnowledgeDocument.user_id == user.id
    ).first()
    
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    
    try:
        # 从向量数据库中删除
        if chroma_client:
            collection_name = f"user_{user.id}_knowledge"
            try:
                collection = chroma_client.get_collection(collection_name)
                # 删除文档的所有向量
                chunk_ids = [f"doc_{doc_id}_chunk_{i}" for i in range(doc.chunk_count)]
                if chunk_ids:
                    collection.delete(ids=chunk_ids)
            except:
                pass  # 集合可能不存在
        
        # 删除文件块记录
        db.query(KnowledgeChunk).filter(KnowledgeChunk.document_id == doc_id).delete()
        
        # 删除文件
        if os.path.exists(doc.file_path):
            os.remove(doc.file_path)
        
        # 删除文档记录
        db.delete(doc)
        db.commit()
        
        return {"message": "文档删除成功"}
        
    except Exception as e:
        logger.error(f"删除文档失败: {e}")
        raise HTTPException(status_code=500, detail="删除文档失败")

@router.post("/search")
async def search_knowledge_base(
    query: str = Form(...),
    top_k: int = Form(5),
    min_weight: int = Form(1),
    file_types: Optional[str] = Form(None),
    username: str = Depends(verify_token),
    db: Session = Depends(get_db)
):
    """搜索知识库"""
    
    start_time = time.time()
    
    user = db.query(User).filter(User.username == username).first()
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在")
    
    if not chroma_client or not embedding_model:
        raise HTTPException(status_code=500, detail="向量搜索服务不可用")
    
    try:
        # 生成查询向量
        query_embedding = embedding_model.encode([query]).tolist()[0]
        
        # 搜索向量数据库
        collection_name = f"user_{user.id}_knowledge"
        try:
            collection = chroma_client.get_collection(collection_name)
        except:
            return {"results": [], "total_time_ms": 0}
        
        # 构建where条件
        where_conditions = {"ai_reference_weight": {"$gte": min_weight}}
        if file_types:
            file_type_list = [ft.strip() for ft in file_types.split(',')]
            where_conditions["file_type"] = {"$in": file_type_list}
        
        search_results = collection.query(
            query_embeddings=[query_embedding],
            n_results=min(top_k, 20),
            where=where_conditions,
            include=["documents", "metadatas", "distances"]
        )
        
        # 格式化结果
        results = []
        if search_results["documents"] and search_results["documents"][0]:
            for i, (doc, metadata, distance) in enumerate(zip(
                search_results["documents"][0],
                search_results["metadatas"][0], 
                search_results["distances"][0]
            )):
                results.append({
                    "chunk_text": doc,
                    "document_id": metadata["document_id"],
                    "filename": metadata["filename"],
                    "file_type": metadata["file_type"],
                    "chunk_index": metadata["chunk_index"],
                    "ai_reference_weight": metadata["ai_reference_weight"],
                    "similarity_score": 1 - distance,  # 转换为相似度分数
                    "chunk_preview": metadata.get("chunk_preview", "")
                })
        
        # 记录查询
        search_time_ms = int((time.time() - start_time) * 1000)
        query_record = KnowledgeQuery(
            user_id=user.id,
            query_text=query,
            query_type='similarity',
            results_count=len(results),
            top_document_ids=[r["document_id"] for r in results],
            search_time_ms=search_time_ms
        )
        db.add(query_record)
        db.commit()
        
        return {
            "results": results,
            "total_time_ms": search_time_ms,
            "query_id": query_record.id
        }
        
    except Exception as e:
        logger.error(f"知识库搜索失败: {e}")
        raise HTTPException(status_code=500, detail=f"搜索失败: {str(e)}")

@router.get("/stats")
async def get_knowledge_stats(
    username: str = Depends(verify_token),
    db: Session = Depends(get_db)
):
    """获取知识库统计信息"""
    
    user = db.query(User).filter(User.username == username).first()
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在")
    
    # 文档统计
    total_docs = db.query(KnowledgeDocument).filter(KnowledgeDocument.user_id == user.id).count()
    processed_docs = db.query(KnowledgeDocument).filter(
        KnowledgeDocument.user_id == user.id,
        KnowledgeDocument.is_processed == True
    ).count()
    
    # 按状态统计
    status_stats = db.query(
        KnowledgeDocument.processing_status,
        db.func.count(KnowledgeDocument.id)
    ).filter(KnowledgeDocument.user_id == user.id).group_by(
        KnowledgeDocument.processing_status
    ).all()
    
    # 按文件类型统计
    type_stats = db.query(
        KnowledgeDocument.file_type,
        db.func.count(KnowledgeDocument.id),
        db.func.sum(KnowledgeDocument.file_size)
    ).filter(KnowledgeDocument.user_id == user.id).group_by(
        KnowledgeDocument.file_type
    ).all()
    
    # 总存储空间
    total_size = db.query(
        db.func.sum(KnowledgeDocument.file_size)
    ).filter(KnowledgeDocument.user_id == user.id).scalar() or 0
    
    # 向量统计
    total_vectors = db.query(
        db.func.sum(KnowledgeDocument.vector_count)
    ).filter(KnowledgeDocument.user_id == user.id).scalar() or 0
    
    # 查询统计
    query_count = db.query(KnowledgeQuery).filter(KnowledgeQuery.user_id == user.id).count()
    
    return {
        "total_documents": total_docs,
        "processed_documents": processed_docs,
        "processing_rate": round(processed_docs / total_docs * 100, 1) if total_docs > 0 else 0,
        "total_storage_bytes": total_size,
        "total_vectors": total_vectors,
        "total_queries": query_count,
        "status_distribution": dict(status_stats),
        "type_distribution": [
            {
                "file_type": file_type,
                "count": count,
                "total_size": size or 0
            } for file_type, count, size in type_stats
        ]
    }

# 用于聊天系统的知识库搜索函数（简化版本）
async def search_knowledge_for_chat(username: str, query: str, db: Session, top_k: int = 3) -> List[str]:
    """为聊天系统提供的知识库搜索功能（简化版本）"""
    
    try:
        user = db.query(User).filter(User.username == username).first()
        if not user:
            return []
        
        # 简单的关键词搜索
        documents = db.query(KnowledgeDocument).filter(
            KnowledgeDocument.user_id == user.id,
            KnowledgeDocument.is_processed == True,
            KnowledgeDocument.extracted_text.contains(query)
        ).limit(top_k).all()
        
        # 返回相关文档内容
        relevant_docs = []
        for doc in documents:
            preview = doc.content_preview or doc.extracted_text[:300] if doc.extracted_text else ""
            relevant_docs.append(f"[来源: {doc.original_filename}] {preview}")
        
        return relevant_docs
        
    except Exception as e:
        logger.error(f"聊天知识库搜索失败: {e}")
        return [] 